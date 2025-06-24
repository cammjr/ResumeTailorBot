import os
from dotenv import load_dotenv
import openai
import gradio as gr
import json
from docx import Document
import tempfile

# Bot label
BOT_NAME = "Resume Tailor 🤵"

# Dependencies:
# pip install openai gradio python-dotenv python-docx

# Load environment variables
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("Please set OPENAI_API_KEY in environment.")
openai.api_key = api_key

# Helper to extract user name from resume
def extract_user_name(resume_text):
    for line in resume_text.splitlines():
        parts = line.strip().split()
        if 2 <= len(parts) <= 4 and all(word.istitle() for word in parts):
            return line.strip()
    return "Candidate"

# Metadata extraction using LLM
def ai_extract_job_meta(job_text):
    """
    Extract company name and job title using OpenAI function calling for strict JSON enforcement.
    """
    # Define function schema for metadata extraction
    functions = [
        {
            "name": "extract_metadata",
            "description": "Extract the company name and job title from a job posting.",
            "parameters": {
                "type": "object",
                "properties": {
                    "company": {"type": "string", "description": "Name of the company offering the job."},
                    "job_title": {"type": "string", "description": "Title of the role being described."}
                },
                "required": ["company", "job_title"]
            }
        }
    ]
    user_msg = {"role": "user", "content": f"Please extract metadata from the following job posting.\n\n{job_text}"}
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "system", "content": "You are an assistant that extracts structured metadata."}, user_msg],
        functions=functions,
        function_call={"name": "extract_metadata"},
        temperature=0
    )
    message = response.choices[0].message
    if getattr(message, "function_call", None):
        args = json.loads(message.function_call.arguments)
        return args.get("company", "Unknown"), args.get("job_title", "Unknown")
    return "Unknown", "Unknown"

# Initialize chat state

def init():
    state = {'step':1, 'resume':'', 'user_name':'', 'company':'', 'job_title':'', 'tailored':''}
    greeting = f"{BOT_NAME}: Hello! I will help tailor your resume with minimal tweaks."
    greeting += "\nStep 1: Paste your original resume text below."
    greeting += "\nType 'New Resume' anytime to restart."
    # Clear all input fields on restart
    return (
        '',  # clear user input box
        [(None, greeting)],  # reset chat history to greeting only
        state,
        gr.update(visible=True, value=''),    # clear and show resume textbox
        gr.update(visible=False, value=''),   # hide and clear job textbox
        gr.update(visible=False, value=None), # hide and clear download choice
        gr.update(visible=False, value=None)  # hide and clear download file
    )

# Main respond function
def respond(input_text, chat, state, resume_text, job_text, download_radio, download_file):
    user_input = input_text.strip()
    # Restart
    if user_input.lower() == 'new resume':
        yield init()
        return

    # Step 1: Receive resume
    if state['step'] == 1:
        if resume_text:
            state['resume'] = resume_text
            state['user_name'] = extract_user_name(resume_text)
            state['step'] = 2
            msg = f"{BOT_NAME}: Got your resume. Now paste the job posting text next."
            yield ('', chat + [(user_input, None), (None, msg)], state,
                   gr.update(visible=False), gr.update(visible=True, value=''),
                   gr.update(visible=False), gr.update(visible=False))
        else:
            msg = f"{BOT_NAME}: Please paste your resume text."
            yield ('', chat + [(None, msg)], state,
                   gr.update(visible=True), gr.update(visible=False),
                   gr.update(visible=False), gr.update(visible=False))
        return

    # Step 2: Receive job posting and tailor
    if state['step'] == 2:
        if job_text:
            busy = f"{BOT_NAME}: Working on your resume..."
            yield ('', chat + [(user_input, None), (None, busy)], state,
                   gr.update(visible=False), gr.update(visible=False),
                   gr.update(visible=False), gr.update(visible=False))

            # Extract metadata
            company, title = ai_extract_job_meta(job_text)
            state['company'], state['job_title'] = company, title
            meta_msgs = [
                (None, f"{BOT_NAME}: Company: {company}"),
                (None, f"{BOT_NAME}: Job Title: {title}"),
                (None, f"{BOT_NAME}: Candidate: {state['user_name']}")
            ]
            yield ('', chat + meta_msgs, state,
                   gr.update(visible=False), gr.update(visible=False),
                   gr.update(visible=False), gr.update(visible=False))

            # Tailor resume
            system_msgs = [
                {'role':'system','content':'You are an expert resume writer.'},
                {'role':'system','content':'First, move the most relevant items to the top of each section, then reword all entries concisely. Do not add new information.'},
                {'role':'system','content':f"Resume:\n{state['resume']}"},
                {'role':'system','content':f"Job Posting:\n{job_text}"}
            ]
            res = openai.chat.completions.create(model='gpt-4o', messages=system_msgs)
            tailored = res.choices[0].message.content
            state['tailored'] = tailored
            yield ('', chat + meta_msgs + [(None, f"{BOT_NAME}: {tailored}")], state,
                   gr.update(visible=False), gr.update(visible=False),
                   gr.update(visible=False), gr.update(visible=False))

            # Explanation
            explain_msgs = [
                {'role':'system','content':'You are an expert resume writer.'},
                {'role':'system','content':'Based on the tailored resume and job posting, here are the items highlighted and why:'},
                {'role':'system','content':f"Tailored Resume:\n{tailored}"},
                {'role':'system','content':f"Job Posting:\n{job_text}"}
            ]
            exp = openai.chat.completions.create(model='gpt-4o', messages=explain_msgs)
            explanation = exp.choices[0].message.content
            # Show download prompt
            prompt = f"{BOT_NAME}: Would you like a Word (.docx) download? (Yes/No)"
            yield ('', chat + meta_msgs + [(None, f"{BOT_NAME}: {tailored}"), (None, f"{BOT_NAME}: {explanation}"), (None, prompt)], state,
                   gr.update(visible=False), gr.update(visible=False),
                   gr.update(visible=True), gr.update(visible=False))
            state['step'] = 3
        else:
            msg = f"{BOT_NAME}: Please paste the job posting text."
            yield ('', chat + [(None, msg)], state,
                   gr.update(visible=False), gr.update(visible=True),
                   gr.update(visible=False), gr.update(visible=False))
        return

    # Step 3: Handle download choice
    if state['step'] == 3 and download_radio:
        if download_radio.lower() == 'yes':
            doc = Document()
            doc.add_heading(state['user_name'], level=1)
            for ln in state['tailored'].split('\n'):
                doc.add_paragraph(ln)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(tmp.name)
            msg = f"{BOT_NAME}: Here is your Word file:"
            yield ('', chat + [(None, msg)], state,
                   gr.update(visible=False), gr.update(visible=False),
                   gr.update(visible=False), gr.update(value=tmp.name, visible=True))
        else:
            msg = f"{BOT_NAME}: Continuing without download. Let me know edits or type New Resume to restart."
            yield ('', chat + [(None, msg)], state,
                   gr.update(visible=False), gr.update(visible=False),
                   gr.update(visible=False), gr.update(visible=False))
        state['step'] = 4
        return

    # Step 4: Ongoing edits
    if state['step'] == 4 and input_text.strip():
        user_edit = input_text.strip()
        yield ('', chat + [(user_edit, None), (None, f"{BOT_NAME}: Applying edits...")], state,
               gr.update(visible=False), gr.update(visible=False),
               gr.update(visible=False), gr.update(visible=False))
        tweak_msgs = [
            {'role':'system', 'content':'You are an expert resume writer.'},
            {'role':'system', 'content':'Tweak the existing tailored resume per user request, maintaining format.'},
            {'role':'system', 'content':f"Tailored Resume:\n{state['tailored']}"}
        ]
        tweak = openai.chat.completions.create(model='gpt-4o', messages=tweak_msgs + [{'role':'user','content':user_edit}])
        state['tailored'] = tweak.choices[0].message.content
        yield ('', chat + [(None, f"{BOT_NAME}: {state['tailored']}")], state,
               gr.update(visible=False), gr.update(visible=False),
               gr.update(visible=False), gr.update(visible=False))
        return

    # Fallback
    yield init()

# UI Definition
with gr.Blocks(title='Resume Tailor 🤵') as demo:
    chatbot = gr.Chatbot(label='Resume Tailor 🤵', type='tuples')
    state = gr.State({})
    resume_text = gr.Textbox(label='Paste your resume text', lines=15, visible=False)
    job_text = gr.Textbox(label='Paste job posting text', lines=10, visible=False)
    download_radio = gr.Radio(['Yes', 'No'], label='Download .docx?', visible=False)
    download_file = gr.File(label='Download .docx', visible=False)
    user_in = gr.Textbox(show_label=False, placeholder='Type here...')

    resume_text.change(respond, [user_in, chatbot, state, resume_text, job_text, download_radio, download_file],
                       [user_in, chatbot, state, resume_text, job_text, download_radio, download_file])
    job_text.change(respond, [user_in, chatbot, state, resume_text, job_text, download_radio, download_file],
                   [user_in, chatbot, state, resume_text, job_text, download_radio, download_file])
    download_radio.change(respond, [user_in, chatbot, state, resume_text, job_text, download_radio, download_file],
                         [user_in, chatbot, state, resume_text, job_text, download_radio, download_file])
    user_in.submit(respond, [user_in, chatbot, state, resume_text, job_text, download_radio, download_file],
                   [user_in, chatbot, state, resume_text, job_text, download_radio, download_file])
    demo.load(init, inputs=None, outputs=[user_in, chatbot, state, resume_text, job_text, download_radio, download_file])
    demo.launch()
